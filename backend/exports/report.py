from __future__ import annotations

from io import BytesIO

from docx import Document

from backend.data.fetcher import build_market_snapshot
from backend.services.analytics import build_options_analytics, build_portfolio_analytics, build_risk_analytics


def build_report_bytes(months: int = 6, risk_free_rate: float = 0.07) -> bytes:
    summary = build_market_snapshot(months=months)
    options_data = build_options_analytics(months=months, risk_free_rate=risk_free_rate)
    portfolio_data = build_portfolio_analytics(months=months, risk_free_rate=risk_free_rate)
    risk_data = build_risk_analytics(months=months)

    document = Document()
    document.add_heading("FRAM Risk Analytics Report", 0)
    document.add_paragraph(
        "This generated report summarizes the current backend scaffold output for liquidity analytics, option pricing, "
        "portfolio hedging, and risk measurement."
    )

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(f"Universe size: {summary['universe_size']} stocks over a {months}-month window.")
    document.add_paragraph(
        f"Liquid bucket size: {len(summary['liquid_bucket'])}. Illiquid bucket size: {len(summary['illiquid_bucket'])}."
    )

    document.add_heading("Part A: Data and Liquidity", level=1)
    for row in summary["liquid_bucket"][:3] + summary["illiquid_bucket"][:3]:
        document.add_paragraph(
            f"{row['symbol']}: avg turnover {row['average_turnover']:.2f}, 20D vol "
            f"{(row['latest_realized_vol_20d'] or 0):.4f}, GARCH forecast {row['garch_forecast_volatility']:.4f}.",
            style="List Bullet",
        )

    document.add_heading("Part B: Option Pricing", level=1)
    for symbol_block in options_data["symbols"]:
        document.add_paragraph(
            f"{symbol_block['symbol']} | Spot {symbol_block['spot']:.2f} | Hist Vol {symbol_block['historical_volatility']:.4f} | "
            f"GARCH Vol {symbol_block['garch_volatility']:.4f}"
        )
        for contract in symbol_block["contracts"][:3]:
            document.add_paragraph(
                f"{contract['label']}: market proxy {contract['market_price_proxy']:.4f}, "
                f"BSM hist {contract['bsm_historical_vol_price']:.4f}, BSM GARCH {contract['bsm_garch_vol_price']:.4f}.",
                style="List Bullet 2",
            )

    document.add_heading("Part C: Portfolio and Hedging", level=1)
    for row in portfolio_data["portfolios"]:
        document.add_paragraph(
            f"{row['symbol']} {row['bucket']} strategy: delta {row['portfolio_greeks']['delta']:.4f}, "
            f"gamma {row['portfolio_greeks']['gamma']:.4f}, vega {row['portfolio_greeks']['vega']:.4f}, "
            f"hedge shares {row['hedge']['raw_shares']:.4f}."
        )

    document.add_heading("Part D: Risk Measurement", level=1)
    for row in risk_data["comparison_table"]:
        document.add_paragraph(
            f"{row['symbol']} ({row['bucket']}, {row['regime']}): VaR95 parametric {row['var_95_parametric']:.2f}, "
            f"VaR99 parametric {row['var_99_parametric']:.2f}, VaR95 GARCH {row['var_95_garch']:.2f}.",
            style="List Bullet",
        )

    document.add_heading("Limitations", level=1)
    document.add_paragraph("Live NSE option-chain market prices are not yet integrated in this scaffold.")
    document.add_paragraph("Shares outstanding and float estimates are currently proxied for turnover ratio calculations.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
